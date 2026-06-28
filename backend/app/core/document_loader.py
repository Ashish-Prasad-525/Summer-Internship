"""
Document Loader - multi-format document parsing.
Supports PDF (pdfplumber), DOCX (python-docx), TXT/MD.
Attaches rich metadata: filename, doc_id, page numbers, char positions.
"""

import logging
import re
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Unified document loader that handles PDF, DOCX, TXT, and MD files.
    Returns a list of LangChain Document objects with rich metadata.
    """

    SUPPORTED = {".pdf", ".docx", ".txt", ".md"}

    def load(self, file_path: str, doc_id: str = "", filename: str = "") -> List[Document]:
        """
        Load a document from disk and return cleaned LangChain Documents.
        Each Document corresponds to one page or logical section.
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED:
            raise ValueError(f"Unsupported file type: {ext}")

        if not filename:
            filename = path.name

        logger.info(f"📂 Loading '{filename}' ({ext})")

        loaders = {
            ".pdf":  self._load_pdf,
            ".docx": self._load_docx,
            ".txt":  self._load_text,
            ".md":   self._load_text,
        }

        docs = loaders[ext](str(path), doc_id=doc_id, filename=filename)
        cleaned = [self._clean_doc(d) for d in docs]
        logger.info(f"✅ Loaded {len(cleaned)} pages/sections from '{filename}'")
        return cleaned

    # ── PDF ───────────────────────────────────────────────────────────────────
    def _load_pdf(self, file_path: str, doc_id: str, filename: str) -> List[Document]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Install pdfplumber: pip install pdfplumber")

        docs = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    continue  # Skip empty pages (scanned images without OCR)

                docs.append(Document(
                    page_content=text,
                    metadata={
                        "doc_id":    doc_id,
                        "filename":  filename,
                        "page":      page_num,
                        "total_pages": len(pdf.pages),
                        "source":    file_path,
                        "file_type": "pdf",
                    },
                ))
        return docs

    # ── DOCX ──────────────────────────────────────────────────────────────────
    def _load_docx(self, file_path: str, doc_id: str, filename: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        docx = DocxDocument(file_path)
        full_text = "\n".join(
            para.text for para in docx.paragraphs if para.text.strip()
        )

        # Split into logical sections (~2000 char blocks for large DOCX)
        docs = []
        section_size = 2000
        for i, start in enumerate(range(0, max(len(full_text), 1), section_size)):
            section = full_text[start: start + section_size]
            if not section.strip():
                continue
            docs.append(Document(
                page_content=section,
                metadata={
                    "doc_id":   doc_id,
                    "filename": filename,
                    "page":     i + 1,
                    "source":   file_path,
                    "file_type": "docx",
                },
            ))
        return docs

    # ── TXT / MD ──────────────────────────────────────────────────────────────
    def _load_text(self, file_path: str, doc_id: str, filename: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        ext = Path(file_path).suffix.lower()
        return [Document(
            page_content=content,
            metadata={
                "doc_id":   doc_id,
                "filename": filename,
                "page":     1,
                "source":   file_path,
                "file_type": ext.lstrip("."),
            },
        )]

    # ── Cleaning ──────────────────────────────────────────────────────────────
    def _clean_doc(self, doc: Document) -> Document:
        """Clean extracted text: normalize whitespace, remove junk characters."""
        text = doc.page_content
        text = re.sub(r"\n{3,}", "\n\n", text)          # Collapse triple+ newlines
        text = re.sub(r"[ \t]{2,}", " ", text)           # Collapse multiple spaces/tabs
        text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text) # Remove non-printable ASCII
        text = text.strip()
        doc.page_content = text
        return doc
